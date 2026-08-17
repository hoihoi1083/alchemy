#!/usr/bin/env python3
"""Generate Alchemy Studio Token Economics & Billing Rules Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Mirror lib/billing/plans.ts + token-costs.ts (Aug 2026)
# Tokens sized for ~75% gross vs provider API cost on Master yearly ($79 / 16,000 ≈ $0.0049375/token).
TOKEN_COGS_USD_PER_1000 = 1.234375
USD_PER_TOKEN = TOKEN_COGS_USD_PER_1000 / 1000
FREE_SIGNUP_GRANT = 500
TOP_UP_TOKENS = 1000
TOP_UP_PRICE_USD = 10

TOKEN_COST = {
    "image": 65,
    "image_ab": 130,
    "campaign": 200,
    "teaching_carousel": 265,
    "storyboard_scene": 65,
    "storyboard_batch": 260,
    "music": 82,
    "voiceover": 13,
    "bgm": 5,
    "plan": 5,
    "inpaint": 41,
    "caption_burn": 8,
}

# Reference-reel video (quality path)
VIDEO_TOKENS_PER_SEC = {
    "480p": 113,
    "720p_fast": 196,
    "720p": 246,
    "1080p": 553,
}

# Single-clip video (studio default)
H3_TOKENS_PER_SEC = {
    "480P": 41,  # $0.05/s · Free cap
    "768P": 65,  # $0.08/s · Standard 720p
    "2K": 106,  # $0.13/s · Pro/Master 1080p
}

HEYGEN_TOKENS_PER_SEC = 82
KLING_5S = 284
KLING_EXTRA_SEC = 57

FREE_PACK_IMAGE = TOKEN_COST["image"]
FREE_PACK_VIDEO_8S = H3_TOKENS_PER_SEC["480P"] * 8  # 328
FREE_PACK_TOTAL = FREE_PACK_IMAGE + FREE_PACK_VIDEO_8S  # 393
FREE_PACK_BUFFER = FREE_SIGNUP_GRANT - FREE_PACK_TOTAL  # 107

STORYBOARD_LANDING = {
    "scenes": 2,
    "clip_sec": 5,
    "total_sec": 10,
    "image_tokens": TOKEN_COST["storyboard_scene"] * 2,  # 130
    "video_tokens": KLING_5S * 2,  # 568
}
STORYBOARD_LANDING["total_tokens"] = (
    STORYBOARD_LANDING["image_tokens"] + STORYBOARD_LANDING["video_tokens"]
)  # 698

PLANS = [
    {
        "name": "Free",
        "tokens": f"{FREE_SIGNUP_GRANT:,} (once at signup)",
        "cogs": f"${FREE_SIGNUP_GRANT * USD_PER_TOKEN:.2f}",
        "price": "$0",
        "margin": "acquisition",
        "notes": "480p video · 1K images · No top-up until paid plan",
    },
    {
        "name": "Standard",
        "tokens": "3,000 / month",
        "cogs": "$3.70",
        "price": "$19.99 / mo",
        "margin": "~81%",
        "notes": "720p video · 1K images · Target plan",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$14.99 / mo yearly",
        "margin": "~75%",
        "notes": "Cheapest Standard token — pricing floor",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$29.99 list",
        "margin": "~88%",
        "notes": "Anchor price",
    },
    {
        "name": "Pro (Most popular)",
        "tokens": "8,000 / month",
        "cogs": "$9.88",
        "price": "$49.99 / mo",
        "margin": "~80%",
        "notes": "1080p video · 1K images",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$39.99 / mo yearly",
        "margin": "~75%",
        "notes": "",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$79.99 list",
        "margin": "~88%",
        "notes": "",
    },
    {
        "name": "Master",
        "tokens": "16,000 / month",
        "cogs": "$19.75",
        "price": "$99.99 / mo",
        "margin": "~80%",
        "notes": "1080p video · 2K images · Pro canvas",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$79.00 / mo yearly",
        "margin": "75%",
        "notes": "Cheapest paid token ($0.00494) — rate used to size all actions",
    },
    {
        "name": "",
        "tokens": "",
        "cogs": "",
        "price": "$159.99 list",
        "margin": "~88%",
        "notes": "",
    },
    {
        "name": "Custom",
        "tokens": "Negotiated",
        "cogs": "—",
        "price": "Contact sales",
        "margin": "—",
        "notes": "Enterprise · 1080p · 2K · Pro canvas",
    },
    {
        "name": "Token top-up",
        "tokens": f"{TOP_UP_TOKENS:,}",
        "cogs": f"${TOP_UP_TOKENS * USD_PER_TOKEN:.2f}",
        "price": f"${TOP_UP_PRICE_USD}",
        "margin": "~88%",
        "notes": "After any paid subscription · $0.01/token",
    },
]


def cogs(tokens: int) -> str:
    return f"~${tokens * USD_PER_TOKEN:.2f}"


def video_8s(res: str) -> int:
    return VIDEO_TOKENS_PER_SEC[res] * 8


def h3_8s(res: str) -> int:
    return H3_TOKENS_PER_SEC[res] * 8


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_fill: str = "1F4E79",
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_heading("Alchemy Studio — Token Economics & Billing Rules", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Generated {date.today().strftime('%Y-%m-%d')} · Internal reference — synced from lib/billing/*"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph(
        f"Anchor: 1,000 tokens ≈ USD {TOKEN_COGS_USD_PER_1000:.2f} operator API cost "
        f"(≈ USD {USD_PER_TOKEN:.4f} / token). Actions are sized so Master yearly "
        f"($79 / 16,000 ≈ $0.00494 / token) still yields ~75% gross vs provider cost. "
        "Retail top-up is $10 / 1,000 ($0.01 / token)."
    )

    # 1. Plans
    doc.add_heading("1. Plan pricing & margin", level=1)
    doc.add_paragraph(
        "Margin assumes the user burns the full monthly token grant. "
        "Yearly Master is the pricing floor (~75% vs provider cost). Monthly and top-up are fatter. "
        "Single-clip video resolution is clamped to the plan cap (Free → 480P, Standard → 768P, Pro/Master → 2K)."
    )
    add_table(
        doc,
        ["Plan", "Tokens", "Full-burn COGS (USD)", "Price charged", "Margin if fully used", "Notes"],
        [[p["name"], p["tokens"], p["cogs"], p["price"], p["margin"], p["notes"]] for p in PLANS],
    )

    # 2. Free pack
    doc.add_heading("2. Free pack promise", level=1)
    doc.add_paragraph(
        f"Free grant: {FREE_SIGNUP_GRANT:,} tokens, once at signup (signupGrantAt). "
        "Product promise: enough for 1 promotional image + 1× 8s 480P video, with buffer for plan/refine."
    )
    add_table(
        doc,
        ["Item", "Tokens", "Approx. COGS"],
        [
            ["Single image (1K)", str(FREE_PACK_IMAGE), cogs(FREE_PACK_IMAGE)],
            ["Video 8s 480P (Free cap)", str(FREE_PACK_VIDEO_8S), cogs(FREE_PACK_VIDEO_8S)],
            ["Pack total", str(FREE_PACK_TOTAL), cogs(FREE_PACK_TOTAL)],
            ["Buffer (plan / refine / retry)", str(FREE_PACK_BUFFER), cogs(FREE_PACK_BUFFER)],
            ["Grant", f"{FREE_SIGNUP_GRANT:,}", f"${FREE_SIGNUP_GRANT * USD_PER_TOKEN:.2f}"],
        ],
    )

    # 3. Token costs
    doc.add_heading("3. Per-action token costs", level=1)
    doc.add_paragraph(
        "Sized so Master yearly (~$0.00494/token) keeps ~75% gross vs provider API cost "
        "(image ~$0.08/1K; single-clip video ~$0.05/s 480P / $0.08/s 768P / $0.13/s 2K; reference-reel published $/s; "
        "presenter ~$0.10/s; stitched clip ~$0.35/5s). "
        "Charge only on successful generation; failed jobs refund via refundTokens()."
    )

    doc.add_heading("Images & packs", level=2)
    add_table(
        doc,
        ["Action", "Est. API cost (USD)", "Tokens"],
        [
            ["Single image (1K)", "~0.08", str(TOKEN_COST["image"])],
            ["A/B (2 images)", "~0.16", str(TOKEN_COST["image_ab"])],
            ["Campaign (3 slides + plan)", "~0.30", str(TOKEN_COST["campaign"])],
            ["Teaching carousel (4 slides + plan)", "~0.40", str(TOKEN_COST["teaching_carousel"])],
            ["Storyboard scene still", "~0.08", str(TOKEN_COST["storyboard_scene"])],
            ["Storyboard batch (4 scenes)", "~0.32", str(TOKEN_COST["storyboard_batch"])],
            ["Inpaint / fill (~1MP)", "~0.05", str(TOKEN_COST["inpaint"])],
            ["Music bed", "~0.10", str(TOKEN_COST["music"])],
            ["Voiceover (short TTS)", "~0.015", str(TOKEN_COST["voiceover"])],
            ["BGM mix (local ffmpeg)", "CPU only", str(TOKEN_COST["bgm"])],
            ["Plan LLM / brief", "~0.01–0.02", str(TOKEN_COST["plan"])],
            ["Caption burn (ffmpeg + R2)", "CPU/storage", str(TOKEN_COST["caption_burn"])],
        ],
    )

    doc.add_heading("Single-clip video (studio default, 5–15s)", level=2)
    add_table(
        doc,
        ["Output", "Tokens / second", "8s total", "Est. API cost / 8s"],
        [
            ["480P (Free · UI 480p)", str(H3_TOKENS_PER_SEC["480P"]), str(h3_8s("480P")), "~$0.40"],
            ["768P (Standard · UI 720p)", str(H3_TOKENS_PER_SEC["768P"]), str(h3_8s("768P")), "~$0.64"],
            ["2K (Pro/Master · UI 1080p)", str(H3_TOKENS_PER_SEC["2K"]), str(h3_8s("2K")), "~$1.04"],
        ],
    )
    doc.add_paragraph(
        "Reference-to-video: first 5 images free; each extra image +65 tokens. "
        "Reference MP4 adds the same tok/s × input seconds (unknown length → assume output duration)."
    )

    doc.add_heading("Reference-reel video (quality path, 4–15s)", level=2)
    add_table(
        doc,
        ["Resolution", "Tokens / second", "8s total", "Est. COGS / 8s"],
        [
            [
                "480p (Free max)",
                str(VIDEO_TOKENS_PER_SEC["480p"]),
                str(video_8s("480p")),
                cogs(video_8s("480p")),
            ],
            [
                "720p fast",
                str(VIDEO_TOKENS_PER_SEC["720p_fast"]),
                str(video_8s("720p_fast")),
                cogs(video_8s("720p_fast")),
            ],
            [
                "720p standard",
                str(VIDEO_TOKENS_PER_SEC["720p"]),
                str(video_8s("720p")),
                cogs(video_8s("720p")),
            ],
            [
                "1080p (Pro/Master max)",
                str(VIDEO_TOKENS_PER_SEC["1080p"]),
                str(video_8s("1080p")),
                cogs(video_8s("1080p")),
            ],
        ],
    )

    doc.add_heading("Presenter video & stitched storyboard", level=2)
    add_table(
        doc,
        ["Action", "Formula", "Example", "Est. COGS"],
        [
            [
                "Talking presenter",
                f"{HEYGEN_TOKENS_PER_SEC} tok/s (bill 4–60s)",
                "15s clip → 1230 tok",
                cogs(HEYGEN_TOKENS_PER_SEC * 15),
            ],
            [
                "Stitched storyboard clip",
                f"{KLING_5S} tok / 5s + {KLING_EXTRA_SEC}/extra s",
                f"5s scene → {KLING_5S} tok",
                cogs(KLING_5S),
            ],
            [
                "Stitched 10s clip",
                f"{KLING_5S} + {KLING_EXTRA_SEC}×5",
                f"10s scene → {KLING_5S + KLING_EXTRA_SEC * 5} tok",
                cogs(KLING_5S + KLING_EXTRA_SEC * 5),
            ],
            [
                "Storyboard landing pack",
                "2 stills + 2×5s stitched (~10s)",
                f"{STORYBOARD_LANDING['total_tokens']} tok/run",
                cogs(STORYBOARD_LANDING["total_tokens"]),
            ],
        ],
    )

    # 4. Approx capacity
    doc.add_heading("4. Approximate monthly capacity (marketing)", level=1)
    doc.add_paragraph(
        "Pricing cards and the landing capacity grid use estimatePricingCardCapacity(). "
        "Each line is an independent maximum prefixed with “Up to”: 1K stills, or 8s 480p video "
        "(lowest studio video resolution, 328 tokens). Mixing formats or using 720p/1080p uses more "
        "tokens per piece. Section 2 still describes the Free pack that fits together "
        "(1 image + 1× 8s 480P)."
    )
    add_table(
        doc,
        ["Plan", "Pricing card (independent maxima at 1K / 8s 480p)"],
        [
            ["Free", "Up to 7 single images · Up to 1 × 8s 480p videos"],
            ["Standard", "Up to 46 single images · Up to 9 × 8s 480p videos"],
            ["Pro", "Up to 123 single images · Up to 24 × 8s 480p videos"],
            ["Master", "Up to 246 single images · Up to 48 × 8s 480p videos"],
        ],
    )
    cap_rows = []
    for plan_name, tokens, video_unit in [
        ("Free", FREE_SIGNUP_GRANT, h3_8s("480P")),
        ("Standard", 3000, h3_8s("480P")),
        ("Pro", 8000, h3_8s("480P")),
        ("Master", 16000, h3_8s("480P")),
    ]:
        cap_rows.append(
            [
                plan_name,
                f"{tokens:,}",
                str(tokens // TOKEN_COST["image"]),
                str(tokens // video_unit),
                str(tokens // STORYBOARD_LANDING["total_tokens"]),
            ]
        )
    add_table(
        doc,
        [
            "Plan",
            "Grant",
            "If grant spent only on 1K images",
            "If grant spent only on 8s video",
            "If grant spent only on ~10s storyboard",
        ],
        cap_rows,
    )

    # 5. Policy
    doc.add_heading("5. Policy rules", level=1)
    policies = [
        "Free grant timing: once at signup only (not monthly refresh). Idempotent via signupGrantAt.",
        "Paid plans: monthly token grant on Stripe invoice.paid (subscription_grant).",
        f"Top-up: ${TOP_UP_PRICE_USD} → {TOP_UP_TOKENS:,} tokens; available only after any paid subscription.",
        "Failed AI jobs: tokens deducted before the generation call; refundTokens() on failure (atomic ledger).",
        "UI: show estimated token cost before Generate; block when balance < cost (402).",
        "Resolution caps: Free 480p/1K · Standard 720p/1K · Pro 1080p/1K · Master 1080p/2K + Pro canvas.",
        "Mongo plan enum: free | standard | pro | master | custom (legacy payg → standard).",
        "Content research (JustOne / Tavily): operator API cost only — not token-charged to users today.",
        "First quick refine per finished image job: still billed as TOKEN_COST.image (65) — no free credit in ledger.",
    ]
    for p in policies:
        doc.add_paragraph(p, style="List Bullet")

    # 6. Implementation status
    doc.add_heading("6. Billing implementation (live)", level=1)
    doc.add_paragraph(
        "Token billing is implemented and wired to generation routes. Stripe Checkout, Customer Portal, "
        "and webhooks sync plan + subscription grants."
    )
    add_table(
        doc,
        ["Component", "Status", "Location"],
        [
            ["Ledger (balance, grant, consume, refund)", "Live", "lib/billing/ledger.ts"],
            ["Pre-check + charge before generate", "Live", "lib/billing/charge.ts"],
            ["Plan definitions & margins", "Live", "lib/billing/plans.ts"],
            ["Per-action costs & estimates", "Live", "lib/billing/token-costs.ts"],
            ["Resolution entitlements", "Live", "lib/billing/entitlements.ts"],
            ["Stripe checkout / portal / webhook", "Live", "app/api/stripe/*, lib/stripe/billing-sync.ts"],
            ["Balance in header", "Live", "GET /api/me"],
            ["Economics unit tests", "Live", "tests/billing-token-costs.test.ts"],
            ["Pricing-card capacity (Free 1+1)", "Live", "lib/billing/pricing-card-capacity.ts"],
        ],
    )

    # 7. Code locations
    doc.add_heading("7. Code locations", level=1)
    locs = [
        "lib/billing/plans.ts — plan definitions, margins, signup grant size",
        "lib/billing/token-costs.ts — TOKEN_COST, videoTokenCost(), FREE_PACK, STORYBOARD_LANDING_PACK",
        "lib/billing/pricing-card-capacity.ts — Free pack vs paid either/or on pricing cards",
        "lib/billing/ledger.ts — assertCanAfford, consumeTokens, grantTokens, refundTokens",
        "lib/billing/charge.ts — chargeTokens (deduct before generate) + refund on failure",
        "lib/billing/entitlements.ts — max video/image resolution per plan",
        "lib/db/types.ts — UserPlan, creditBalance, signupGrantAt, Stripe fields",
        "app/api/generate-image/route.ts, app/api/generate/route.ts — primary deduct paths",
        "app/api/generate-kling-storyboard/route.ts — storyboard video billing",
        "tests/billing-token-costs.test.ts — economics unit tests",
    ]
    for loc in locs:
        doc.add_paragraph(loc, style="List Bullet")

    doc.add_paragraph(
        "Re-generate this document: python scripts/generate-token-economics-docx.py "
        "(requires python-docx in .venv-docx)."
    )

    return doc


def main() -> None:
    out_paths = [
        Path("/Users/michaelng/Desktop/Alchemy-Studio-Token-Economics.docx"),
        Path("/Users/michaelng/Desktop/alchemy-studio/docs/Alchemy-Studio-Token-Economics.docx"),
    ]
    doc = build_doc()
    for p in out_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        print(f"Wrote {p}")


if __name__ == "__main__":
    main()
