#!/usr/bin/env python3
"""Generate Alchemy Studio — complete token charging reference (Word)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Mirror lib/billing/token-costs.ts + plans.ts (Aug 2026)
TOKEN_COGS_USD_PER_1000 = 1.234375
USD_PER_TOKEN = TOKEN_COGS_USD_PER_1000 / 1000
FREE_SIGNUP_GRANT = 300
TOP_UP_TOKENS = 1000
TOP_UP_PRICE_USD = 10

TOKEN_COST = {
    "image": 65,
    "image_ab": 130,
    "campaign": 200,
    "plan": 5,
    "storyboard_scene": 65,
    "music": 82,
    "voiceover": 13,
    "bgm": 5,
    "inpaint": 41,
    "caption_burn": 8,
    "smart_layers_detect": 8,
    "smart_layers_matte": 5,
    "smart_layers_heal": 3,
}

H3_TOKENS_PER_SEC = {
    "480P": 41,
    "768P": 65,
    "2K": 106,
    "4K": 130,
}

VIDEO_TOKENS_PER_SEC = {
    "480p": 113,
    "720p_fast": 196,
    "720p": 246,
    "1080p": 553,
}

HEYGEN_TOKENS_PER_SEC = 82
KLING_5S = 284
KLING_10S = 284 + 57 * 5  # 569


def teaching_carousel(n: int) -> int:
    n = max(4, min(6, n))
    return TOKEN_COST["plan"] + TOKEN_COST["image"] * n


def storyboard_stills(n: int, passes: int = 1) -> int:
    return TOKEN_COST["storyboard_scene"] * n * passes


def kling_storyboard(n: int, clip_sec: int = 5) -> int:
    per = KLING_5S if clip_sec <= 5 else KLING_10S
    return per * n


def h3_8s(res: str) -> int:
    return H3_TOKENS_PER_SEC[res] * 8


def seedance_8s(res: str, fast: bool = False) -> int:
    key = "720p_fast" if res == "720p" and fast else res
    return VIDEO_TOKENS_PER_SEC[key] * 8


def add_title(doc: Document, text: str) -> None:
    t = doc.add_heading(text, level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "Alchemy Studio — Token Charging Guide")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Generated {date.today().isoformat()} · alchemy-studio")
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph(
        "This document lists every place Alchemy Studio charges tokens, how each cost is "
        "calculated, what is free, and how research billing works. Source of truth: "
        "lib/billing/token-costs.ts and app/api/* routes that call chargeTokens()."
    )

    # 1. How billing works
    doc.add_heading("1. How billing works", level=1)
    for item in [
        "Tokens are deducted before the AI job runs (chargeTokens). Failed jobs are refunded (refundTokens).",
        "Balance is shown in the header and on Account → Receipts & token history.",
        "Each charge stores meta.kind (e.g. research_reel, storyboard, image) for itemized history labels.",
        "Insufficient balance returns HTTP 402; the UI shows estimated cost before Generate.",
        f"Free signup grant: {FREE_SIGNUP_GRANT:,} tokens once (not monthly).",
        f"Top-up (paid plans only): ${TOP_UP_PRICE_USD} → {TOP_UP_TOKENS:,} tokens.",
        "Tokens expire 6 months after grant; oldest tokens spent first (FIFO).",
        f"Operator COGS reference: 1,000 tokens ≈ ${TOKEN_COGS_USD_PER_1000:.4f} fal/API cost (~${USD_PER_TOKEN:.6f}/token).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # 2. Flat catalog
    doc.add_heading("2. Flat token catalog", level=1)
    add_table(
        doc,
        ["Action / unit", "Tokens", "Notes"],
        [
            ["Single 1K image (Nano Banana)", str(TOKEN_COST["image"]), "Base image unit — 65 tok per still"],
            ["A/B two variants", str(TOKEN_COST["image_ab"]), "2 × 65"],
            ["Campaign set (3 images)", str(TOKEN_COST["campaign"]), "3 × 65 + 5 plan"],
            ["Teaching carousel (4 slides)", str(teaching_carousel(4)), "5 plan + 4 × 65"],
            ["Teaching carousel (5 slides)", str(teaching_carousel(5)), "Default slide count"],
            ["Teaching carousel (6 slides)", str(teaching_carousel(6)), "Max slides"],
            ["Storyboard still (1 scene)", str(TOKEN_COST["storyboard_scene"]), "Same rate as single image"],
            ["Storyboard still regen (1 cell)", str(TOKEN_COST["storyboard_scene"]), "Per-cell regen"],
            ["LLM plan / brief step", str(TOKEN_COST["plan"]), "DeepSeek planning calls"],
            ["Music generation", str(TOKEN_COST["music"]), "MiniMax music on fal"],
            ["Voiceover preview / TTS", str(TOKEN_COST["voiceover"]), "Per voice clip"],
            ["BGM mix (ffmpeg)", str(TOKEN_COST["bgm"]), "Local processing fee"],
            ["Caption burn-in", str(TOKEN_COST["caption_burn"]), "ffmpeg + storage"],
            ["Image inpaint / erase (FLUX Fill)", f"{TOKEN_COST['inpaint']} × MP", "Per rounded megapixel"],
            ["Smart layers — detect", str(TOKEN_COST["smart_layers_detect"]), "Florence OCR + detect"],
            ["Smart layers — matte", str(TOKEN_COST["smart_layers_matte"]), "BiRefNet per layer"],
            ["Smart layers — heal (local)", str(TOKEN_COST["smart_layers_heal"]), "Ring-fill, no FLUX"],
            ["HeyGen presenter", f"{HEYGEN_TOKENS_PER_SEC}/sec", "4–60 sec billed"],
            ["Kling I2V clip", f"{KLING_5S} (5s) · {KLING_10S} (10s)", "Per storyboard scene clip"],
        ],
    )

    # 3. Research
    doc.add_heading("3. Research — how it charges", level=1)
    doc.add_heading("3a. FREE to users (0 tokens)", level=2)
    add_table(
        doc,
        ["API route", "What it does", "Operator cost (you pay)"],
        [
            [
                "POST /api/research-content-angles",
                "Keyword search on 小红书 / Instagram → post cards + DeepSeek angles",
                "JustOne ~¥0.08–0.12/search + DeepSeek ~$0.001–0.01",
            ],
            [
                "POST /api/research-direct-post",
                "Paste a post URL → analyze that post + angles",
                "JustOne detail fetch + DeepSeek",
            ],
            [
                "POST /api/plan-storyboard",
                "LLM storyboard plan from brief (no images yet)",
                "DeepSeek only (~$0.01–0.02) — no chargeTokens",
            ],
            [
                "POST /api/analyze-reference",
                "Reference image vision brief",
                "fal vision — no chargeTokens today",
            ],
            [
                "POST /api/research-resolve-video",
                "Resolve platform video URL to MP4",
                "JustOne / CDN fetch — no chargeTokens",
            ],
            [
                "POST /api/research-post-video",
                "Proxy-download research reel bytes",
                "CDN bandwidth — no chargeTokens",
            ],
        ],
    )
    doc.add_paragraph(
        "Important: Platform keyword research in Step 1 is intentionally free to encourage discovery. "
        "JustOne and DeepSeek costs are operator-only."
    )

    doc.add_heading("3b. PAID research (charges tokens)", level=2)
    add_table(
        doc,
        ["API route", "Tokens", "Formula", "Billing label"],
        [
            [
                "POST /api/analyze-research-reel",
                "70 or 75",
                f"65 vision + 5 plan + (5 storyboard plan if enabled)",
                "Research reel analysis (+ storyboard plan)",
            ],
            [
                "POST /api/refine-research-video-script",
                str(TOKEN_COST["plan"]),
                "DeepSeek script refine after reel analysis",
                "Research video script refine",
            ],
        ],
    )
    doc.add_paragraph(
        "Reel analysis is NOT the same as one image: 75 tokens vs 65 for a single image. "
        "The 65-token portion covers fal vision (same rate as image-class); the extra 5–10 tokens "
        "cover DeepSeek planning. Default includes auto storyboard plan (+5)."
    )

    doc.add_heading("3c. Example — Ronald Kwok session (Aug 27, 2026)", level=2)
    add_table(
        doc,
        ["Time (HKT)", "Delta", "Balance after", "What happened"],
        [
            ["3:13 PM", "+300", "300", "Free signup grant"],
            ["3:18 PM", "−75", "225", "Reel analysis #1 (+ storyboard plan)"],
            ["3:23 PM", "−130", "95", "Storyboard images — 2 scenes × 65"],
            ["5:22 PM", "−75", "20", "Reel analysis #2 (+ storyboard plan)"],
        ],
    )

    # 4. Images
    doc.add_heading("4. Image generation", level=1)
    add_table(
        doc,
        ["API route", "Tokens", "When / formula", "meta.kind"],
        [
            [
                "POST /api/generate-image",
                "65–260",
                "1–4 images: 65×count; A/B=130; via imageOutputMode",
                "image",
            ],
            [
                "POST /api/generate-image (refine)",
                "65–260",
                "Quick fix / logo refine / region refine — per numImages",
                "image (mode: refine)",
            ],
            [
                "POST /api/generate-campaign",
                "200 or 65",
                "Full set=200; single slide regen=65",
                "campaign / image",
            ],
            [
                "POST /api/generate-teaching-carousel",
                "265–395 or 65",
                "5+65×slides (4–6); single slide regen=65",
                "teaching_carousel / image",
            ],
            [
                "POST /api/generate-cinematic-scenes",
                "65 × scenes",
                "Cinematic stitch stills",
                "cinematic_scenes",
            ],
            [
                "POST /api/inpaint-image",
                "41 × MP",
                "FLUX Fill generative erase",
                "inpaint",
            ],
        ],
    )
    doc.add_paragraph("Wizard output modes → token estimate (sidebar):")
    add_table(
        doc,
        ["Output mode", "Typical tokens"],
        [
            ["Single", "65"],
            ["A/B two versions", "130"],
            ["Campaign (3 posts)", "200"],
            ["Teaching carousel 4/5/6 slides", f"{teaching_carousel(4)} / {teaching_carousel(5)} / {teaching_carousel(6)}"],
        ],
    )

    # 5. Storyboard
    doc.add_heading("5. Storyboard", level=1)
    doc.add_paragraph("Storyboard is two separate charges: stills first, then video animation.")
    add_table(
        doc,
        ["Step", "API route", "Tokens", "Formula"],
        [
            [
                "Storyboard stills",
                "POST /api/generate-storyboard-images",
                "65 × scenes × passes",
                "passes=2 if brand logo Mode A (Nano edit per scene)",
            ],
            [
                "Storyboard video (primary)",
                "POST /api/generate-kling-storyboard",
                "H3 or Seedance",
                "Tries MiniMax H3 or Seedance R2V first based on balance/engine",
            ],
            [
                "Storyboard video (Kling fallback)",
                "POST /api/generate-kling-storyboard",
                f"{KLING_5S}–{KLING_10S} × scenes",
                "Stitched Kling I2V per scene if H3/Seedance unavailable",
            ],
        ],
    )
    doc.add_paragraph("Storyboard still examples:")
    add_table(
        doc,
        ["Scenes", "Tokens (1 pass)", "Tokens (logo Mode A, 2 passes)"],
        [
            ["2", str(storyboard_stills(2)), str(storyboard_stills(2, 2))],
            ["4", str(storyboard_stills(4)), str(storyboard_stills(4, 2))],
            ["6", str(storyboard_stills(6)), str(storyboard_stills(6, 2))],
        ],
    )
    doc.add_paragraph("Full storyboard reel example (2 scenes × 5s Kling + stills):")
    add_table(
        doc,
        ["Component", "Tokens"],
        [
            ["Stills (2 scenes)", str(storyboard_stills(2))],
            ["Kling video (2 × 5s)", str(kling_storyboard(2, 5))],
            ["Total", str(storyboard_stills(2) + kling_storyboard(2, 5))],
        ],
    )

    # 6. Video
    doc.add_heading("6. Video generation", level=1)
    doc.add_heading("6a. MiniMax H3 (studio default single-clip path)", level=2)
    add_table(
        doc,
        ["Resolution", "Tokens/sec", "8s clip", "Plan cap"],
        [
            ["480P", "41", str(h3_8s("480P")), "Free"],
            ["768P (720p UI)", "65", str(h3_8s("768P")), "Standard / Light"],
            ["2K (1080p UI)", "106", str(h3_8s("2K")), "Pro / Master"],
            ["4K", "130", str(H3_TOKENS_PER_SEC['4K'] * 8), "Master"],
        ],
    )
    doc.add_paragraph(
        "Extra reference images after the first 5: +65 tokens each. "
        "Reference video input: billed at same $/s × input seconds (capped 15s)."
    )

    doc.add_heading("6b. Seedance (quality / reel path)", level=2)
    add_table(
        doc,
        ["Resolution", "Tokens/sec", "8s clip"],
        [
            ["480p", "113", str(seedance_8s("480p"))],
            ["720p fast", "196", str(seedance_8s("720p", fast=True))],
            ["720p", "246", str(seedance_8s("720p"))],
            ["1080p", "553", str(seedance_8s("1080p"))],
        ],
    )

    doc.add_heading("6c. Video API routes", level=2)
    add_table(
        doc,
        ["API route", "Engine", "Tokens", "meta.kind"],
        [
            ["POST /api/generate", "Seedance", "113–553 × sec (4–15s)", "video"],
            ["POST /api/generate-minimax-h3", "MiniMax H3", "41–130 × sec", "minimax_h3"],
            ["POST /api/generate-kling-storyboard", "H3 / Seedance / Kling", "See §5", "minimax_h3 / video / kling_storyboard_fallback"],
            ["POST /api/generate-digital-presenter", "HeyGen + optional TTS", "82×sec + 13 voice", "digital_presenter"],
        ],
    )

    # 7. Audio & captions
    doc.add_heading("7. Audio, captions & post-production", level=1)
    add_table(
        doc,
        ["API route", "Tokens", "Notes", "meta.kind"],
        [
            ["POST /api/generate-music", str(TOKEN_COST["music"]), "AI background music", "music"],
            ["POST /api/preview-script-voice", str(TOKEN_COST["voiceover"]), "TTS preview", "voiceover"],
            ["POST /api/dub-script-voice", "13 × lines", "Multi-line dub", "voiceover_dub"],
            ["POST /api/add-bgm", str(TOKEN_COST["bgm"]), "Mix uploaded BGM", "bgm"],
            ["POST /api/plan-caption-voice", str(TOKEN_COST["plan"]), "Caption timing plan", "caption_plan"],
            ["POST /api/expand-caption-voice", str(TOKEN_COST["plan"]), "Expand caption text", "caption_expand"],
            ["POST /api/expand-spoken-captions", str(TOKEN_COST["plan"]), "Spoken caption expand", "caption_expand_spoken"],
            ["POST /api/burn-script-captions", str(TOKEN_COST["caption_burn"]), "Burn captions on video", "caption_burn"],
            ["POST /api/burn-visual-captions", str(TOKEN_COST["caption_burn"]), "Visual caption burn", "caption_burn"],
            ["POST /api/finish-blockbuster", str(TOKEN_COST["caption_burn"]), "Final caption pass", "finish-blockbuster"],
            [
                "POST /api/postprocess",
                "8–26+",
                "8 base + 5 ASR + 5 rewrite + 13 dub (optional stack)",
                "postprocess",
            ],
        ],
    )

    # 8. Edit-image toolkit
    doc.add_heading("8. Edit-image / Pro toolkit", level=1)
    add_table(
        doc,
        ["API route", "Tokens", "Tool", "meta.kind"],
        [
            ["POST /api/decompose-image-layers", "8", "Smart layers detect", "smart_layers_detect"],
            ["POST /api/layer-matte", "5", "Layer matte / cutout", "smart_layers_matte"],
            ["POST /api/layer-heal", "3 or 41×MP", "Local heal or FLUX erase", "smart_layers_heal"],
            ["POST /api/inpaint-image", "41×MP", "Generative inpaint", "inpaint"],
        ],
    )

    # 9. Complete charge map
    doc.add_heading("9. Complete API charge map (all chargeTokens routes)", level=1)
    routes = [
        ("POST /api/analyze-research-reel", "Research", "70–75", "research_reel"),
        ("POST /api/refine-research-video-script", "Research", "5", "refine_research_video_script"),
        ("POST /api/generate-image", "Image", "65–260", "image"),
        ("POST /api/generate-campaign", "Image", "65 or 200", "campaign / image"),
        ("POST /api/generate-teaching-carousel", "Image", "65 or 265–395", "teaching_carousel / image"),
        ("POST /api/generate-storyboard-images", "Storyboard", "65×scenes×passes", "storyboard"),
        ("POST /api/generate-cinematic-scenes", "Image", "65×scenes", "cinematic_scenes"),
        ("POST /api/generate", "Video", "Seedance per-sec", "video"),
        ("POST /api/generate-minimax-h3", "Video", "H3 per-sec", "minimax_h3"),
        ("POST /api/generate-kling-storyboard", "Video", "H3/Seedance/Kling", "minimax_h3 / video / kling_storyboard_fallback"),
        ("POST /api/generate-digital-presenter", "Video", "HeyGen per-sec + voice", "digital_presenter"),
        ("POST /api/generate-music", "Audio", "82", "music"),
        ("POST /api/preview-script-voice", "Audio", "13", "voiceover"),
        ("POST /api/dub-script-voice", "Audio", "13×lines", "voiceover_dub"),
        ("POST /api/add-bgm", "Audio", "5", "bgm"),
        ("POST /api/plan-caption-voice", "Captions", "5", "caption_plan"),
        ("POST /api/expand-caption-voice", "Captions", "5", "caption_expand"),
        ("POST /api/expand-spoken-captions", "Captions", "5", "caption_expand_spoken"),
        ("POST /api/burn-script-captions", "Captions", "8", "caption_burn"),
        ("POST /api/burn-visual-captions", "Captions", "8", "caption_burn"),
        ("POST /api/finish-blockbuster", "Captions", "8", "finish-blockbuster"),
        ("POST /api/postprocess", "Post", "variable", "postprocess"),
        ("POST /api/inpaint-image", "Edit", "41×MP", "inpaint"),
        ("POST /api/decompose-image-layers", "Edit", "8", "smart_layers_detect"),
        ("POST /api/layer-matte", "Edit", "5", "smart_layers_matte"),
        ("POST /api/layer-heal", "Edit", "3 or 41×MP", "smart_layers_heal"),
    ]
    add_table(doc, ["Route", "Category", "Tokens", "meta.kind"], routes)

    # 10. NOT charged
    doc.add_heading("10. Routes that do NOT charge tokens", level=1)
    free_routes = [
        "/api/research-content-angles — platform keyword research",
        "/api/research-direct-post — direct post URL research",
        "/api/plan-storyboard — LLM storyboard planning only",
        "/api/analyze-reference — reference image brief",
        "/api/research-resolve-video — video URL resolution",
        "/api/research-post-video — video proxy download",
        "/api/stitch-videos — ffmpeg stitch (no chargeTokens)",
        "/api/trim-video — ffmpeg trim",
        "/api/compose — compositor layout (no fal generation charge)",
        "/api/archive-studio-image — library save",
        "/api/studio-download — export download",
        "/api/burn-image-canvas — canvas export",
        "/api/burn-image-text — text overlay export",
        "/api/batch-export-images — batch zip export",
        "Stripe grants — signup_grant, subscription_grant, trial_bonus, topup (credit, not debit)",
    ]
    for r in free_routes:
        doc.add_paragraph(r, style="List Bullet")

    # 11. Plans
    doc.add_heading("11. Plan token grants (reference)", level=1)
    add_table(
        doc,
        ["Plan", "Tokens", "Video cap", "Image cap"],
        [
            ["Free (signup once)", "300", "480p", "1K"],
            ["Light", "3,000/mo", "480p", "1K"],
            ["Standard", "8,000/mo", "720p", "1K"],
            ["Pro", "16,000/mo", "1080p", "1K"],
            ["Master", "28,000/mo", "1080p", "2K + Pro canvas"],
            ["Custom (Enterprise)", "40,000/mo", "1080p", "2K + Pro canvas"],
        ],
    )

    # 12. Code locations
    doc.add_heading("12. Code locations", level=1)
    for loc in [
        "lib/billing/token-costs.ts — TOKEN_COST, estimates, video/image formulas",
        "lib/billing/charge.ts — chargeTokens(), refundTokens(), imageTokenCostFromRequest()",
        "lib/billing/ledger.ts — Mongo credit_transactions (reason + meta)",
        "lib/billing/transaction-label.ts — Account history labels from meta.kind",
        "components/AccountPageClient.tsx — Receipts & token history UI",
        "docs/XHS_NOTE_SEARCH_SETUP.md — JustOne operator cost notes",
    ]:
        doc.add_paragraph(loc, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph("Re-generate: python3 scripts/generate-token-charging-guide-docx.py")
    p.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    return doc


def main() -> None:
    out = Path("/Users/michaelng/Desktop/Alchemy-Studio-Token-Charging-Guide.docx")
    doc = build_doc()
    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
