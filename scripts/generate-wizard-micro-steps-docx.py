#!/usr/bin/env python3
"""Convert docs/WIZARD_MICRO_STEPS.md to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "WIZARD_MICRO_STEPS.md"
OUT_PATH = ROOT / "docs" / "WIZARD_MICRO_STEPS.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_md_table(doc: Document, header_line: str, rows: list[str]) -> None:
    headers = [c.strip() for c in header_line.strip("|").split("|")]
    body: list[list[str]] = []
    for row in rows:
        if re.match(r"^\|[-:\s|]+\|$", row):
            continue
        cells = [c.strip() for c in row.strip("|").split("|")]
        if len(cells) == len(headers):
            body.append(cells)
    if not headers or not body:
        return
    table = doc.add_table(rows=1 + len(body), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1F4E79")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row in enumerate(body):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 41, 59)


def add_rich_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def convert_md_to_docx(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    i = 0
    in_code = False
    code_buf: list[str] = []
    table_header: str | None = None
    table_rows: list[str] = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            in_code = True
            i += 1
            continue

        if line.startswith("|"):
            if table_header is None:
                table_header = line
                table_rows = []
            else:
                table_rows.append(line)
            if i + 1 >= len(lines) or not lines[i + 1].startswith("|"):
                add_md_table(doc, table_header, table_rows)
                table_header = None
                table_rows = []
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            add_rich_paragraph(doc, re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        else:
            add_rich_paragraph(doc, line.strip())

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> int:
    md = MD_PATH
    out = OUT_PATH
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not md.is_file():
        print(f"Missing {md}", file=sys.stderr)
        return 1
    convert_md_to_docx(md, out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
