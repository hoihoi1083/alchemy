#!/usr/bin/env python3
"""Generate Alchemy Studio usage & pricing summary Word document."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# FX for display (approximate — document as reference only)
USD_TO_HKD = 7.80
USD_TO_CNY = 7.20


def usd(usd_low: float, usd_high: float | None = None) -> str:
    if usd_high is None or usd_high == usd_low:
        hkd = usd_low * USD_TO_HKD
        cny = usd_low * USD_TO_CNY
        return f"USD {usd_low:.2f}  |  HKD {hkd:.0f}  |  CNY {cny:.2f}"
    hkd_l, hkd_h = usd_low * USD_TO_HKD, usd_high * USD_TO_HKD
    cny_l, cny_h = usd_low * USD_TO_CNY, usd_high * USD_TO_CNY
    return (
        f"USD {usd_low:.2f}–{usd_high:.2f}  |  "
        f"HKD {hkd_l:.0f}–{hkd_h:.0f}  |  "
        f"CNY {cny_l:.2f}–{cny_h:.2f}"
    )


def cny(cny_low: float, cny_high: float | None = None) -> str:
    if cny_high is None or cny_high == cny_low:
        hkd = cny_low * (USD_TO_HKD / USD_TO_CNY)
        usd_val = cny_low / USD_TO_CNY
        return f"CNY {cny_low:.2f}  |  HKD {hkd:.2f}  |  USD {usd_val:.3f}"
    hkd_l = cny_low * (USD_TO_HKD / USD_TO_CNY)
    hkd_h = cny_high * (USD_TO_HKD / USD_TO_CNY)
    usd_l, usd_h = cny_low / USD_TO_CNY, cny_high / USD_TO_CNY
    return (
        f"CNY {cny_low:.2f}–{cny_high:.2f}  |  "
        f"HKD {hkd_l:.2f}–{hkd_h:.2f}  |  "
        f"USD {usd_l:.3f}–{usd_h:.3f}"
    )


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], header_fill: str = "1F4E79") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = doc.add_heading("Alchemy Studio — Usage & Pricing Summary", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"Prepared: {date.today().strftime('%d %B %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph(
        "This document summarises what each product action costs you (API pass-through) "
        "and how you should charge end users. Prices are estimates from internal benchmarks "
        "and vendor dashboards — verify before publishing a price list."
    )

    # --- Charging model ---
    doc.add_heading("1. How charging works today", level=1)
    doc.add_paragraph(
        "User billing is not live yet. The app records usage in the database for analytics only; "
        "there is no Stripe checkout and no automatic deduction from user wallets. "
        "All generation today is paid by your API keys (FAL, DeepSeek, Just One API, etc.)."
    )
    add_table(
        doc,
        ["Topic", "Current status", "Recommended approach"],
        [
            [
                "Billing",
                "Not implemented — creditBalance field exists but stays at 0",
                "Charge fixed service packages (HKD), not raw API passthrough",
            ],
            [
                "Regenerate",
                "Every new image / scene / video = new API run",
                "Show cost hint in UI; warn before regenerate",
            ],
            [
                "Quick minor edit",
                "1 free refine per finished job, then counts as a full image",
                "Include 1 tweak in package price; extra edits = add-on",
            ],
            [
                "Research search",
                "Billed to your Just One / Tavily account per click",
                "Bundle 10–20 searches/month in a plan or charge per research run",
            ],
            [
                "Local post-process",
                "Caption burn, stitch, BGM mix — server ffmpeg only",
                "Free to user (no external API cost)",
            ],
        ],
    )

    doc.add_heading("Suggested retail prices (what to charge SMB customers)", level=2)
    add_table(
        doc,
        ["Package", "Suggested retail (HKD)", "Your typical API cost (HKD)"],
        [
            ["Single social image (with or without reference)", "HKD 50 – 150", "HKD 0.3 – 0.6"],
            ["Campaign set — 3 linked posts", "HKD 200 – 400", "HKD 2 – 3"],
            ["Teaching carousel — 4 cards", "HKD 250 – 500", "HKD 3 – 4"],
            ["8-second Reel (image + video + optional VO/BGM)", "HKD 300 – 800", "HKD 12 – 16"],
            ["Full funnel (research + carousel + Reel)", "HKD 800 – 1,500", "HKD 20 – 40"],
        ],
    )
    doc.add_paragraph(
        "Target margin: roughly 10–30× API cost for self-serve, or fixed project fees for agency work. "
        "Do not quote API wholesale prices to customers."
    )

    # --- Search ---
    doc.add_heading("2. Content search (market research)", level=1)
    doc.add_paragraph(
        "Users search RedNote / Instagram / TikTok / Facebook by keyword, or paste a direct post URL. "
        "The system fetches trending posts, then an LLM synthesises angles and can auto-fill the wizard."
    )
    add_table(
        doc,
        ["Action", "What happens", "Your API cost"],
        [
            [
                "One platform research click",
                "1 live search on chosen platform + 1 LLM synthesis of angles",
                cny(0.08, 0.12) + " + " + usd(0.001, 0.01),
            ],
            [
                "Pin / open a specific post",
                "Extra detail API call (needed for direct URL or video post)",
                cny(0.08, 0.12),
            ],
            [
                "Tavily fallback (if primary search fails)",
                "Web search backup when configured",
                usd(0.02, 0.05),
            ],
            [
                "Typical UI session",
                "1 search + synthesis; +1 detail if user picks a post",
                cny(0.08, 0.24) + " + LLM",
            ],
        ],
    )
    doc.add_paragraph(
        "Charging note: One research click ≈ CNY 0.10–0.12 (HKD 0.11–0.13) plus negligible LLM cost. "
        "At scale, CNY 36 (~USD 5) buys roughly 350 searches. "
        "Charge users per research run (e.g. HKD 5–15) or include a monthly search allowance in a plan."
    )

    # --- Image ---
    doc.add_heading("3. Image generation", level=1)
    doc.add_paragraph(
        "Social ad images for IG / Facebook / RedNote. Includes single image, A/B variants, "
        "3-slide campaign sets, and 4-card teaching carousels. Planning step uses a small LLM call."
    )
    add_table(
        doc,
        ["Action", "What happens", "Your API cost"],
        [
            ["Single promotional image", "1 image generation or edit", usd(0.04, 0.08)],
            ["A/B two versions", "2 images from same settings", usd(0.08, 0.16)],
            [
                "Campaign set (3 slides)",
                "LLM plans 3 linked posts, then 3 images",
                usd(0.24, 0.30) + " (incl. plan)",
            ],
            [
                "Teaching carousel (4 slides)",
                "LLM plans 4 edu cards, then 4 images",
                usd(0.32, 0.40) + " (incl. plan)",
            ],
            [
                "Storyboard scene still (per scene)",
                "1 image per scene before video",
                usd(0.06),
            ],
            [
                "Storyboard planning",
                "LLM writes scene list and motion notes",
                usd(0.002),
            ],
            [
                "Auto second frame (video prep)",
                "Optional extra still before video",
                usd(0.04, 0.08),
            ],
            [
                "Quick minor edit (1st per job)",
                "Small change on existing image",
                "HKD 0 (1 free credit per job)",
            ],
            [
                "Quick minor edit (after free credit)",
                "Same as new image generation",
                usd(0.04, 0.08),
            ],
            [
                "Brand / concept analyse (Setup)",
                "LLM (+ optional vision note) fills wizard fields",
                usd(0.002, 0.01),
            ],
        ],
    )
    doc.add_paragraph(
        "Charging note: Quote per deliverable (1 image, 3-pack, 4-pack), not per API call. "
        "Campaign and carousel are ~3× and ~4× the cost of a single image respectively."
    )

    # --- Reference ---
    doc.add_heading("4. Using a reference (style / layout / product)", level=1)
    doc.add_paragraph(
        "Users can upload a reference ad image (or pin one from research). "
        "The app analyses layout, colours, and typography, then generates new creative "
        "with the user's headline and product — not a pixel copy of the reference."
    )
    add_table(
        doc,
        ["Reference mode", "What happens", "Your API cost"],
        [
            [
                "Analyse reference image",
                "Vision reads layout, palette, typography → creative brief",
                usd(0.01),
            ],
            [
                "Style-only reference",
                "Reference pixels sent with prompt; new layout + user copy",
                usd(0.04, 0.08) + " (same as 1 image; analyse once)",
            ],
            [
                "Layout transfer (reference + product photo)",
                "Reference + product both sent; mirror layout, swap subject",
                usd(0.08),
            ],
            [
                "Product photo only (no reference ad)",
                "Product image as hero; brand fields style the ad",
                usd(0.04, 0.08),
            ],
            [
                "Multi-slide reference carousel (vision)",
                "Per-slide layout DNA from several reference images",
                usd(0.01) + " per analyse batch",
            ],
            [
                "Reference reel video (content research)",
                "Download MP4 → analyse shots → plan storyboard scenes",
                usd(0.01, 0.02) + " LLM plan",
            ],
        ],
    )
    doc.add_paragraph(
        "Charging note: Reference analyse is cheap (≈ HKD 0.08). "
        "The main cost is still each generated image or video. "
        "Style-only and layout-transfer do not add a separate 'reference fee' in the API — "
        "price your packages to include one reference upload per job."
    )

    # --- Video ---
    doc.add_heading("5. Video generation", level=1)
    doc.add_paragraph(
        "Short-form Reels / RedNote video from a keyframe image, storyboard scenes, or reference clip. "
        "Optional AI music, Cantonese voiceover, and burned-in captions (local ffmpeg — no extra API)."
    )
    add_table(
        doc,
        ["Action", "What happens", "Your API cost"],
        [
            ["4-second clip, draft quality (480p fast)", "Image-to-video or text-to-video", usd(0.95)],
            ["6–8 second clip, draft quality (480p fast)", "Typical single Reel", usd(1.45, 1.90)],
            ["10-second clip, higher quality (720p standard)", "Premium tier", usd(3.00) + "+"],
            ["Text-to-video (no product photo)", "Prompt-only video", usd(0.95) + " (4s fast)"],
            ["Reference-to-video (uploaded MP4)", "Motion matched to reference clip", usd(1.45, 1.90)],
            [
                "Storyboard Reel",
                "N scene images + 1 video (or N videos if cinematic stitch)",
                f"N × {usd(0.06)} images + video(s)",
            ],
            [
                "Cinematic stitch (e.g. 3 × 8s)",
                "3 keyframes + 3 video clips + local stitch",
                usd(4.35, 5.70) + " images",
            ],
            ["AI voiceover (short Cantonese / Mandarin)", "TTS after video", usd(0.01)],
            ["AI background music", "Generated bed under video", "Varies — budget USD 0.05–0.15"],
            ["Caption burn / BGM mix / splice", "Server-side ffmpeg", "HKD 0"],
        ],
    )
    doc.add_paragraph(
        "Charging note: Video is the highest-cost step. "
        "An 8s draft Reel costs you roughly HKD 11–15 in API fees. "
        "Retail at HKD 300–800 per Reel depending on polish (VO, music, captions, revisions). "
        "Fast/draft mode is ~⅓ the cost of a 24s multi-scene stitch."
    )

    # --- Combined workflows ---
    doc.add_heading("6. Combined workflows (your cost checklist)", level=1)
    add_table(
        doc,
        ["User journey", "Your total API cost (approx.)"],
        [
            ["Research → single image with reference", "CNY 0.10 + USD 0.05–0.10  |  HKD 1 – 2"],
            ["Campaign 3-pack with reference", "USD 0.25–0.35  |  HKD 2 – 3"],
            ["Teaching carousel 4-pack with reference", "USD 0.35–0.45  |  HKD 3 – 4"],
            ["Image → 8s Reel (draft)", "USD 1.55–2.00  |  HKD 12 – 16"],
            ["Image → 8s Reel + VO + BGM", "USD 1.60–2.15  |  HKD 13 – 17"],
            ["Research → storyboard (3 scenes) → video", "USD 2.00–2.50  |  HKD 16 – 20"],
        ],
    )

    # --- FX footnote ---
    doc.add_heading("7. Exchange rates used in this document", level=1)
    doc.add_paragraph(
        f"USD 1 = HKD {USD_TO_HKD:.2f}  |  USD 1 = CNY {USD_TO_CNY:.2f}  |  "
        f"CNY 1 ≈ HKD {USD_TO_HKD / USD_TO_CNY:.2f}. "
        "Update these before publishing customer-facing prices."
    )
    doc.add_paragraph(
        "Sources: internal pipeline smoke benchmarks, PRODUCT_EVALUATION.md, XHS_NOTE_SEARCH_SETUP.md, "
        "and vendor dashboard pricing (Just One API, fal.ai, DeepSeek). "
        "API prices change — re-check quarterly."
    )

    return doc


def main() -> None:
    out_paths = [
        Path("/Users/michaelng/Desktop/alchemy-studio/docs/Alchemy-Studio-Usage-Pricing-Summary.docx"),
        Path("/Users/michaelng/Desktop/Alchemy-Studio-Usage-Pricing-Summary.docx"),
    ]
    doc = build_doc()
    for p in out_paths:
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        print(f"Wrote {p}")


if __name__ == "__main__":
    main()
